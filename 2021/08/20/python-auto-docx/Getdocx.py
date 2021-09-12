import docx 
import os
from PIL import Image
from datetime import datetime
from datetime import timedelta
from datetime import timezone
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from easydict import EasyDict  
import re
import pandas as pd

if __name__ == "__main__":
    conf = EasyDict()  
    conf.webUrl='https://www.douyin.com/'
    conf.webSource='官方网站'
    conf.dangeStr=['低风险','中风险','高风险']
    conf.appName='摩尔庄园'
    conf.appSource='应用市场'
    conf.imgSuffix=['png','jpg','jpeg']
    page_num=1
    SHA_TZ = timezone(
        timedelta(hours=8),
        name='Asia Beijing',
    )
    utc_now = datetime.utcnow().replace(tzinfo=timezone.utc)
    beijing_now = utc_now.astimezone(SHA_TZ)
    # print(beijing_now, beijing_now.tzname())
    # print(beijing_now)

    resTime=str(beijing_now)+'  '+str(beijing_now.tzname())
    # print(resTime)
    resTime= re.sub(r'\.[0-9]*', ' ',resTime)
    # print(resTime)
    # print(beijing_now.date(), beijing_now.tzname())
    #打开demo.docx文档，返回一个Document对象，它有paragraphs属性，时Paragraph对象的列表。
    doc = docx.Document()
    header = doc.sections[0].header
    paragraph = header.paragraphs[0]
    paragraph.add_run('云账户@www.yunzhanghu.com')
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paraObj0=doc.add_paragraph('多模态风险评估','Title')
    paraObj0.alignment=WD_ALIGN_PARAGRAPH.CENTER
    # paraObj0.runs[0].add_break()
    paraObj01=doc.add_paragraph('云账户')
    paraObj01.alignment=WD_ALIGN_PARAGRAPH.CENTER

    paraObj1=doc.add_paragraph(resTime)
    paraObj1.alignment=WD_ALIGN_PARAGRAPH.CENTER

    #一、评估对象
    doc.add_heading('一、评估对象',0)
    table = doc.add_table(rows=1, cols=3,style='Table Grid')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '主体类别'
    hdr_cells[1].text = '主体信息'
    hdr_cells[2].text = '来源'

    row_cells = table.add_row().cells
    row_cells[0].text = 'Web网页'
    row_cells[1].text = conf.webUrl
    row_cells[2].text = conf.webSource
    
    row_cells = table.add_row().cells
    row_cells[0].text = 'App应用'
    row_cells[1].text = conf.appName
    row_cells[2].text = conf.appSource
    paraObj10=doc.add_paragraph(' ')
    paraObj10.runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)

    # footer = doc.sections[0].footer # 获取第一个节的页脚
    # paragraph = footer.paragraphs[0] # 获取页脚的第一个段落
    # paragraph.add_run(str(page_num))#添加内容
    # paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER#居中

    #二、Web端评估信息
    doc.add_heading('二、Web端评估信息',0)


    # paraObj2 = doc.add_paragraph('This is a second paragraph.')
    # paraObj2.add_run('This text is being added to the second paragraph.真的')


    # paraObj3 = doc.add_paragraph('This is a yet another paragraph.')
    # paraObj3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  

    df = pd.read_excel('./web.xlsx')
    print('web imgs number is',len(df.index.values))
    table = doc.add_table(rows=1, cols=3,style='Table Grid')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '图片名称'
    hdr_cells[1].text = '危险程度'
    hdr_cells[2].text = '置信度'
    for i in range(len(df.index.values)):
        img=df.loc[i].values
        pic_name,danger,confident=img
        print(pic_name,danger,confident)
        suffix=pic_name.split('.')[-1]
       
        if suffix in conf.imgSuffix:
            print(pic_name)
            
            row_cells = table.add_row().cells
            row_cells[0].text = str(pic_name)
            row_cells[1].text = conf.dangeStr[danger]
            row_cells[2].text = str(confident)
            
    # # footer = doc.sections[1].footer # 获取第一个节的页脚
    # paragraph = footer.paragraphs[1] # 获取页脚的第一个段落
    # page_num=page_num+1
    # paragraph.add_run(str(page_num))#添加内容
    # paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER#居中

    paraObj=doc.add_paragraph(' ')
    paraObj.runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)




    for i in range(len(df.index.values)):
        img=df.loc[i].values
        pic_name,danger,confident=img
        print(pic_name,danger,confident)
        suffix=pic_name.split('.')[-1]
       
        if suffix in conf.imgSuffix:
            print(pic_name)
            imgPath='./webPic/'+pic_name
            image = Image.open(imgPath)
            size=image.size
            print(size)
            w,h=size
            delta=h/w
            if delta>1:
                trueW=10
                trueH=delta*trueW
                if trueH>15:
                    trueH=15
                    trueW=trueH/delta
            elif delta==1:
                trueW=10
                trueH=delta*trueW
            else:
                trueH=10
                trueW=trueH/delta
                if trueW>15:
                    trueW=15
                    trueH=trueW*delta

            paragraph101 = doc.add_paragraph()
            paragraph101.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER     
            run = paragraph101.add_run("")
            run.add_picture(imgPath,width=docx.shared.Cm(trueW),height=docx.shared.Cm(trueH))
            
            table = doc.add_table(rows=1, cols=3,style='Table Grid')
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '图片名称'
            hdr_cells[1].text = '危险程度'
            hdr_cells[2].text = '置信度'

            row_cells = table.add_row().cells
            row_cells[0].text = str(pic_name)
            row_cells[1].text = conf.dangeStr[danger]
            row_cells[2].text = str(confident)
            
            
            
            paraObj10=doc.add_paragraph(' ')
            paraObj10.runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
    # for file in files :
    #     suffix=file.split('.')[-1]
    #     imgSuffix=['png','jpg','jpeg']
    #     if suffix in imgSuffix:
    #         image = Image.open(file)
    #         size=image.size
    #         print(size)
    #         w,h=size
    #         delta=h/w
    #         doc.add_picture(file,width=docx.shared.Cm(10),height=docx.shared.Cm(10*delta))




    doc.add_heading('三、App端评估信息',0)
    df = pd.read_excel('./app.xlsx')
    print('app imgs number is',len(df.index.values))
    table = doc.add_table(rows=1, cols=3,style='Table Grid')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '图片名称'
    hdr_cells[1].text = '危险程度'
    hdr_cells[2].text = '置信度'
    for i in range(len(df.index.values)):
        img=df.loc[i].values
        pic_name,danger,confident=img
        print(pic_name,danger,confident)
        suffix=pic_name.split('.')[-1]
       
        if suffix in conf.imgSuffix:
            print(pic_name)
            
            row_cells = table.add_row().cells
            row_cells[0].text = str(pic_name)
            row_cells[1].text = conf.dangeStr[danger]
            row_cells[2].text = str(confident)
            
            
    paraObj=doc.add_paragraph(' ')
    paraObj.runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)




    for i in range(len(df.index.values)):
        img=df.loc[i].values
        pic_name,danger,confident=img
        print(pic_name,danger,confident)
        suffix=pic_name.split('.')[-1]
       
        if suffix in conf.imgSuffix:
            print(pic_name)
            imgPath='./appPic/'+pic_name
            image = Image.open(imgPath)
            size=image.size
            print(size)
            w,h=size
            delta=h/w
            if delta>1:
                trueW=10
                trueH=delta*trueW
                if trueH>15:
                    trueH=15
                    trueW=trueH/delta
            elif delta==1:
                trueW=10
                trueH=delta*trueW
            else:
                trueH=10
                trueW=trueH/delta
                if trueW>15:
                    trueW=15
                    trueH=trueW*delta

            paragraph101 = doc.add_paragraph()
            paragraph101.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER     
            run = paragraph101.add_run("")
            run.add_picture(imgPath,width=docx.shared.Cm(trueW),height=docx.shared.Cm(trueH))
            
            table = doc.add_table(rows=1, cols=3,style='Table Grid')
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '图片名称'
            hdr_cells[1].text = '危险程度'
            hdr_cells[2].text = '置信度'

            row_cells = table.add_row().cells
            row_cells[0].text = str(pic_name)
            row_cells[1].text = conf.dangeStr[danger]
            row_cells[2].text = str(confident)
            
            
            
            # paraObj10=doc.add_paragraph(' ')
            # paraObj10.runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
 
    # doc.add_heading('Header 0',0)
    # doc.add_heading('Header 1',1)
    # doc.add_heading('Header 2',2)
    # doc.add_heading('Header 3',3)
    # doc.add_heading('Header 4',4)


    doc.save('风险评测结果.docx')
